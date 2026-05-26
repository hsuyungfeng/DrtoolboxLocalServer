#!/usr/bin/env python3
"""將行銷文案匯出為 .docx 格式（支援嵌入圖片）"""

import argparse
import os
import sys

try:
    from docx import Document
    from docx.shared import Pt, Cm, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("請先安裝 python-docx：pip install python-docx 或 uv pip install python-docx")
    sys.exit(1)


def create_marketing_docx(
    title: str,
    content: str,
    output_path: str,
    platform: str = "",
    images: list = None,
):
    """建立行銷文案 Word 文件，可嵌入圖片"""
    doc = Document()

    # 設定頁面邊距
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # 標題
    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 平台標註
    if platform:
        platform_para = doc.add_paragraph()
        platform_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = platform_para.add_run(f"【{platform}】")
        run.font.size = Pt(12)

    # 分隔線
    doc.add_paragraph("─" * 40)

    # 主視覺圖（第一張放在文案前面）
    if images and len(images) > 0 and os.path.exists(images[0]):
        img_para = doc.add_paragraph()
        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = img_para.add_run()
        run.add_picture(images[0], width=Inches(4.5))
        doc.add_paragraph("")

    # 文案內容
    for line in content.split("\n"):
        para = doc.add_paragraph(line)
        para.paragraph_format.space_after = Pt(6)

    # 其餘圖片放在文案後面
    if images and len(images) > 1:
        doc.add_paragraph("")
        doc.add_paragraph("─" * 40)
        caption_para = doc.add_paragraph("【搭配圖片素材】")
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in caption_para.runs:
            run.font.size = Pt(11)
            run.bold = True
        doc.add_paragraph("")

        for img_path in images[1:]:
            if os.path.exists(img_path):
                img_para = doc.add_paragraph()
                img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = img_para.add_run()
                run.add_picture(img_path, width=Inches(4.0))
                # 圖片檔名作為說明
                fname = os.path.basename(img_path)
                cap = doc.add_paragraph(f"▲ {fname}")
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in cap.runs:
                    r.font.size = Pt(8)
                    r.font.italic = True
                doc.add_paragraph("")

    # 頁尾：診所資訊
    doc.add_paragraph("─" * 40)
    footer_text = "緻妍外科診所 Zhiyan Surgical Clinic\n台中市太平區新福路208號\n電話：04-2395-0960"
    footer_para = doc.add_paragraph(footer_text)
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer_para.runs:
        run.font.size = Pt(9)

    doc.save(output_path)
    print(f"文案已匯出至：{output_path}")


def main():
    parser = argparse.ArgumentParser(description="行銷文案匯出工具")
    parser.add_argument("--title", required=True, help="文案標題")
    parser.add_argument("--content", required=True, help="文案內容（支援 \\n 換行）")
    parser.add_argument("--output", required=True, help="輸出檔案路徑（.docx）")
    parser.add_argument("--platform", default="", help="平台標註（如 Instagram / Facebook）")
    parser.add_argument("--images", nargs="*", default=[], help="圖片路徑（可多張，第一張為主視覺）")
    args = parser.parse_args()

    content = args.content.replace("\\n", "\n")
    create_marketing_docx(args.title, content, args.output, args.platform, args.images)


if __name__ == "__main__":
    main()

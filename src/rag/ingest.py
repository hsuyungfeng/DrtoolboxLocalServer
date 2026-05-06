"""
DocumentIngestor - RAG document ingestion pipeline.

This module provides a DocumentIngestor class for ingesting and indexing
medical documents (PDF, Word, plain text) into Chroma vector database.

Features:
- Multi-format document parsing (PDF, DOCX, TXT)
- Text chunking with overlap
- Semantic embedding with metadata tracking
- Progress reporting and error handling
"""

import os
import time
import logging
import hashlib
from typing import List, Optional, Generator, Dict, Any
from dataclasses import dataclass, field
from pathlib import Path

# Document parsing
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Vector database
try:
    import chromadb
    from chromadb.config import Settings
    CHROMA_AVAILABLE = True
except ImportError:
    CHROMA_AVAILABLE = False


logger = logging.getLogger(__name__)


@dataclass
class ChunkResult:
    """Result of document chunking."""
    chunk_id: str
    text: str
    source: str
    page: Optional[int]
    section: Optional[str]
    index: int
    tokens: int


@dataclass
class IngestResult:
    """Result of document ingestion."""
    source: str
    chunks: int
    embedded: bool
    collection: str
    duration_ms: float
    errors: List[str] = field(default_factory=list)


class DocumentIngestor:
    """
    Document ingestion pipeline for RAG.
    
    Handles parsing, chunking, and embedding of medical documents
    into Chroma vector database.
    
    Attributes:
        chroma_dir: Path to Chroma persistence directory
        collection_name: Name of Chroma collection
        chunk_size: Tokens per chunk
        chunk_overlap: Overlap between chunks
    """
    
    def __init__(
        self,
        chroma_dir: str = "data/rag/chroma/",
        collection_name: str = "medical_documents",
        chunk_size: int = 512,
        chunk_overlap: int = 50,
        embedding_model: str = "sentence-transformers/all-MiniLM-L6-v2",
    ):
        """
        Initialize DocumentIngestor.
        
        Args:
            chroma_dir: Path to Chroma persistence directory
            collection_name: Name of Chroma collection
            chunk_size: Tokens per chunk
            chunk_overlap: Overlap between chunks
            embedding_model: Sentence-transformers model name
        """
        self.chroma_dir = chroma_dir
        self.collection_name = collection_name
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.embedding_model = embedding_model
        
        self.client = None
        self.collection = None
        self._documents: List[str] = []
        self._metadatas: List[dict] = []
        self._ids: List[str] = []
        
        # Supported formats
        self.supported_formats = {'.txt', '.pdf', '.docx'}
        
        logger.info(f"DocumentIngestor initialized: collection={collection_name}")
    
    def _init_chroma(self) -> bool:
        """Initialize Chroma client."""
        if not CHROMA_AVAILABLE:
            logger.error("chromadb not installed")
            return False
        
        try:
            # Use existing Chroma instance if available
            os.makedirs(self.chroma_dir, exist_ok=True)
            
            self.client = chromadb.PersistentClient(
                path=self.chroma_dir,
                settings=Settings(
                    anonymized_telemetry=False,
                    allow_reset=True,
                )
            )
            
            # Get or create collection
            self.collection = self.client.get_or_create_collection(
                name=self.collection_name,
                metadata={"description": "Medical documents for RAG"}
            )
            
            logger.info(f"Chroma collection ready: {self.collection_name}")
            return True
            
        except Exception as e:
            logger.error(f"Failed to initialize Chroma: {e}")
            return False
    
    def parse_document(self, file_path: str) -> str:
        """
        Parse document and extract text.
        
        Args:
            file_path: Path to document file
            
        Returns:
            Extracted text content
            
        Raises:
            ValueError: If file format not supported
            RuntimeError: If parsing fails
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Document not found: {file_path}")
        
        ext = Path(file_path).suffix.lower()
        
        if ext not in self.supported_formats:
            raise ValueError(f"Unsupported format: {ext}")
        
        text = ""
        
        try:
            if ext == '.txt':
                with open(file_path, 'r', encoding='utf-8') as f:
                    text = f.read()
                    
            elif ext == '.pdf' and PDF_AVAILABLE:
                text = self._parse_pdf(file_path)
                
            elif ext == '.docx' and DOCX_AVAILABLE:
                text = self._parse_docx(file_path)
            
            return text
            
        except Exception as e:
            raise RuntimeError(f"Failed to parse {file_path}: {e}")
    
    def _parse_pdf(self, file_path: str) -> str:
        """Parse PDF file."""
        text_parts = []
        
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            
            for page_num, page in enumerate(reader.pages, start=1):
                page_text = page.extract_text()
                
                if page_text:
                    text_parts.append(f"[Page {page_num}]\n{page_text}")
        
        return "\n\n".join(text_parts)
    
    def _parse_docx(self, file_path: str) -> str:
        """Parse Word document."""
        doc = docx.Document(file_path)
        text_parts = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)
        
        return "\n\n".join(text_parts)
    
    def chunk_text(
        self,
        text: str,
        source: str,
        page_start: int = 0,
    ) -> List[ChunkResult]:
        """
        Split text into chunks.
        
        Args:
            text: Full text content
            source: Source file path
            page_start: Starting page number (for metadata)
            
        Returns:
            List of ChunkResult objects
        """
        chunks = []
        
        # Simple token-based chunking (approximate: ~4 chars per token)
        chars_per_token = 4
        chunk_chars = self.chunk_size * chars_per_token
        overlap_chars = self.chunk_overlap * chars_per_token
        
        text_len = len(text)
        
        if text_len <= chunk_chars:
            # Single chunk
            chunk_id = self._generate_chunk_id(source, 0)
            chunks.append(ChunkResult(
                chunk_id=chunk_id,
                text=text,
                source=source,
                page=None,
                section=None,
                index=0,
                tokens=text_len // chars_per_token,
            ))
            return chunks
        
        # Multiple chunks with overlap
        position = 0
        index = 0
        
        while position < text_len:
            end = min(position + chunk_chars, text_len)
            
            # Try to break at sentence boundary
            if end < text_len:
                for boundary in ['. ', '.\n', '!\n', '?\n', '\n\n']:
                    last_boundary = text.rfind(boundry, position, end)
                    if last_boundary > position + chunk_chars // 2:
                        end = last_boundary + len(boundary)
                        break
            
            chunk_text = text[position:end].strip()
            
            if chunk_text:
                chunk_id = self._generate_chunk_id(source, index)
                chunk_tokens = len(chunk_text) // chars_per_token
                
                # Try to detect section from first line
                section = None
                first_line = chunk_text.split('\n')[0].strip()
                if first_line and len(first_line) < 100:
                    section = first_line
                
                chunks.append(ChunkResult(
                    chunk_id=chunk_id,
                    text=chunk_text,
                    source=source,
                    page=page_start + index if page_start else None,
                    section=section,
                    index=index,
                    tokens=chunk_tokens,
                ))
            
            # Move position with overlap
            position = end - overlap_chars
            index += 1
        
        return chunks
    
    def _generate_chunk_id(self, source: str, index: int) -> str:
        """Generate unique chunk ID."""
        content = f"{source}:{index}"
        hash_suffix = hashlib.md5(content.encode()).hexdigest()[:8]
        return f"chunk_{Path(source).stem}_{index}_{hash_suffix}"
    
    def ingest(
        self,
        file_path: str,
        metadata: Optional[Dict[str, Any]] = None,
    ) -> IngestResult:
        """
        Ingest document into vector database.
        
        Args:
            file_path: Path to document file
            metadata: Additional metadata for document
            
        Returns:
            IngestResult with ingestion status
        """
        start_time = time.time()
        errors = []
        
        # Initialize Chroma if not ready
        if self.client is None:
            if not self._init_chroma():
                raise RuntimeError("Failed to initialize Chroma")
        
        try:
            # Parse document
            text = self.parse_document(file_path)
            
            if not text.strip():
                errors.append(f"No text extracted from {file_path}")
            
            # Chunk text
            chunks = self.chunk_text(text, file_path)
            
            # Prepare for embedding
            documents = [chunk.text for chunk in chunks]
            ids = [chunk.chunk_id for chunk in chunks]
            
            metadatas = []
            base_meta = {
                "source": os.path.basename(file_path),
                "full_path": file_path,
                "ingested_at": time.time(),
            }
            if metadata:
                base_meta.update(metadata)
            
            for chunk in chunks:
                meta = base_meta.copy()
                meta.update({
                    "page": chunk.page,
                    "section": chunk.section,
                    "chunk_index": chunk.index,
                    "tokens": chunk.tokens,
                })
                metadatas.append(meta)
            
            # Add to Chroma
            self.collection.add(
                documents=documents,
                ids=ids,
                metadatas=metadatas,
            )
            
            duration_ms = (time.time() - start_time) * 1000
            
            result = IngestResult(
                source=file_path,
                chunks=len(chunks),
                embedded=True,
                collection=self.collection_name,
                duration_ms=duration_ms,
                errors=errors,
            )
            
            logger.info(
                f"Ingested {file_path}: {len(chunks)} chunks in {duration_ms:.0f}ms"
            )
            
            return result
            
        except Exception as e:
            logger.error(f"Ingestion failed for {file_path}: {e}")
            raise
    
    def ingest_batch(
        self,
        file_paths: List[str],
        metadata: Optional[Dict[str, Any]] = None,
    ) -> List[IngestResult]:
        """
        Ingest multiple documents.
        
        Args:
            file_paths: List of document paths
            metadata: Additional metadata
            
        Returns:
            List of IngestResult objects
        """
        results = []
        
        for file_path in file_paths:
            try:
                result = self.ingest(file_path, metadata)
                results.append(result)
            except Exception as e:
                logger.error(f"Failed to ingest {file_path}: {e}")
                results.append(IngestResult(
                    source=file_path,
                    chunks=0,
                    embedded=False,
                    collection=self.collection_name,
                    duration_ms=0,
                    errors=[str(e)],
                ))
        
        return results
    
    def search(
        self,
        query: str,
        n_results: int = 5,
        where: Optional[dict] = None,
    ) -> List[Dict[str, Any]]:
        """
        Search documents by semantic similarity.
        
        Args:
            query: Search query text
            n_results: Number of results to return
            where: Optional metadata filter
            
        Returns:
            List of matching documents with scores
        """
        if self.collection is None:
            raise RuntimeError("Collection not initialized")
        
        results = self.collection.query(
            query_texts=[query],
            n_results=n_results,
            where=where,
        )
        
        # Format results
        formatted = []
        
        if results.get('documents') and results['documents'][0]:
            for i, doc in enumerate(results['documents'][0]):
                formatted.append({
                    "text": doc,
                    "id": results['ids'][0][i],
                    "distance": results['distances'][0][i],
                    "metadata": results['metadatas'][0][i],
                })
        
        return formatted
    
    def get_collection_info(self) -> dict:
        """Get collection information."""
        if self.collection is None:
            return {"initialized": False}
        
        return {
            "initialized": True,
            "name": self.collection_name,
            "count": self.collection.count(),
            "chroma_dir": self.chroma_dir,
        }
    
    def delete_collection(self, confirm: bool = False):
        """Delete collection and all embedded documents."""
        if not confirm:
            raise ValueError("Must set confirm=True to delete collection")
        
        if self.client:
            self.client.delete_collection(self.collection_name)
            logger.info(f"Collection deleted: {self.collection_name}")
            self.collection = None
    
    def __enter__(self):
        """Context manager entry."""
        self._init_chroma()
        return self
    
    def __exit__(self, exc_type, exc_val, exc_tb):
        """Context manager exit."""
        pass